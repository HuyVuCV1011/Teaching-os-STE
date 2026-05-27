import os
import json
import subprocess
import tempfile
import unittest
from docx import Document
import openpyxl

class TestParseMaterial(unittest.TestCase):
    def test_parse_docx(self):
        # Create temp DOCX
        doc = Document()
        doc.add_heading('Testing Document Parsing', level=1)
        
        p = doc.add_paragraph()
        p.add_run('This is a ').bold = True
        p.add_run('bold').bold = True
        p.add_run(' and ')
        p.add_run('italic').italic = True
        p.add_run(' sentence.')
        
        doc.add_paragraph('Bullet 1', style='List Bullet')
        doc.add_paragraph('Bullet 2', style='List Bullet')
        
        table = doc.add_table(rows=2, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Header 1'
        hdr_cells[1].text = 'Header 2'
        
        row_cells = table.rows[1].cells
        row_cells[0].text = 'Val 1'
        row_cells[1].text = 'Val 2'
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            tmp_name = tmp.name
            
        try:
            # Run CLI parser
            cmd = [os.path.join(os.path.dirname(__file__), '../.venv/bin/python'), 
                   os.path.join(os.path.dirname(__file__), '../scripts/parse_material.py'), 
                   tmp_name]
            
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            data = json.loads(result.stdout)
            
            self.assertIn("viewer_artifact", data)
            self.assertIn("extracted_text", data)
            
            artifact = data["viewer_artifact"]
            self.assertEqual(artifact["type"], "docx")
            self.assertIn("<h1>Testing Document Parsing</h1>", artifact["viewer_html"])
            self.assertIn("<strong>bold</strong>", artifact["viewer_html"])
            self.assertIn("<em>italic</em>", artifact["viewer_html"])
            self.assertIn("Header 1", artifact["viewer_html"])
            self.assertIn("Val 1", artifact["viewer_html"])
            
            self.assertIn("# Testing Document Parsing", artifact["viewer_markdown"])
            self.assertIn("**bold**", artifact["viewer_markdown"])
            self.assertIn("*italic*", artifact["viewer_markdown"])
            
            self.assertIn("Testing Document Parsing", data["extracted_text"])
            self.assertIn("Bullet 1", data["extracted_text"])
            
        finally:
            if os.path.exists(tmp_name):
                os.remove(tmp_name)

    def test_parse_csv(self):
        # Create temp CSV
        with tempfile.NamedTemporaryFile(suffix='.csv', delete=False, mode='w', encoding='utf-8') as tmp:
            tmp.write("Name,Age,Role\n")
            tmp.write("Alice,30,Engineer\n")
            tmp.write("Bob,25,Designer\n")
            tmp_name = tmp.name
            
        try:
            # Run CLI parser
            cmd = [os.path.join(os.path.dirname(__file__), '../.venv/bin/python'), 
                   os.path.join(os.path.dirname(__file__), '../scripts/parse_material.py'), 
                   tmp_name]
            
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            data = json.loads(result.stdout)
            
            self.assertIn("viewer_artifact", data)
            self.assertIn("extracted_text", data)
            
            artifact = data["viewer_artifact"]
            self.assertEqual(artifact["type"], "tabular")
            self.assertEqual(artifact["headers"], ["Name", "Age", "Role"])
            self.assertEqual(artifact["rows"], [["Alice", "30", "Engineer"], ["Bob", "25", "Designer"]])
            self.assertEqual(artifact["row_count"], 2)
            self.assertEqual(artifact["col_count"], 3)
            
            self.assertIn("Alice, 30, Engineer", data["extracted_text"])
            
        finally:
            if os.path.exists(tmp_name):
                os.remove(tmp_name)

    def test_parse_xlsx(self):
        # Create temp XLSX
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        ws.append(["ID", "Item", "Price"])
        ws.append(["1", "Laptop", "1000"])
        ws.append(["2", "Mouse", "20"])
        
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            wb.save(tmp.name)
            tmp_name = tmp.name
            
        try:
            # Run CLI parser
            cmd = [os.path.join(os.path.dirname(__file__), '../.venv/bin/python'), 
                   os.path.join(os.path.dirname(__file__), '../scripts/parse_material.py'), 
                   tmp_name]
            
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            data = json.loads(result.stdout)
            
            self.assertIn("viewer_artifact", data)
            self.assertIn("extracted_text", data)
            
            artifact = data["viewer_artifact"]
            self.assertEqual(artifact["type"], "tabular")
            self.assertEqual(artifact["headers"], ["ID", "Item", "Price"])
            self.assertEqual(artifact["rows"], [["1", "Laptop", "1000"], ["2", "Mouse", "20"]])
            self.assertEqual(artifact["row_count"], 2)
            self.assertEqual(artifact["col_count"], 3)
            self.assertEqual(artifact["sheet_names"], ["Test Sheet"])
            
            self.assertIn("Laptop, 1000", data["extracted_text"])
            
        finally:
            if os.path.exists(tmp_name):
                os.remove(tmp_name)

if __name__ == '__main__':
    unittest.main()
